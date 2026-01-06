"""
Document Extractor Service.
Extracts text from PDF, DOCX, XLSX files.
Memory-only - no disk writes.
"""
import io
import zipfile
from typing import Optional, Tuple

from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook
import pandas as pd

from database import supabase

# Document classification keywords
CLASSIFICATION_KEYWORDS = {
    "AVIS": ["avis de consultation", "aoon", "aooi", "avis d'appel"],
    "RC": ["règlement de consultation", "reglement de consultation", "r.c."],
    "CPS": ["cahier des prescriptions spéciales", "c.p.s.", "cahier des charges"],
    "ANNEXE": ["annexe", "additif", "avenant"],
}


class DocumentExtractor:
    """Extract text from various document formats."""

    def __init__(self):
        self.ocr_engine = None  # Lazy load PaddleOCR

    async def extract_and_store(
        self, tender_id: str, filename: str, file_bytes: io.BytesIO
    ):
        """Extract documents from a tender package and store them."""
        file_bytes.seek(0)

        # Check if it's a ZIP file
        if filename.lower().endswith(".zip"):
            await self._process_zip(tender_id, file_bytes)
        else:
            # Single file
            content, method, pages = self._extract_single(filename, file_bytes)
            doc_type = self._classify_document(content)
            await self._store_document(
                tender_id, filename, doc_type, content, method, pages
            )

    async def _process_zip(self, tender_id: str, file_bytes: io.BytesIO):
        """Process a ZIP file containing multiple documents."""
        try:
            with zipfile.ZipFile(file_bytes) as zf:
                for name in zf.namelist():
                    if name.startswith("__MACOSX") or name.endswith("/"):
                        continue

                    with zf.open(name) as f:
                        inner_bytes = io.BytesIO(f.read())
                        content, method, pages = self._extract_single(name, inner_bytes)
                        doc_type = self._classify_document(content)
                        await self._store_document(
                            tender_id, name, doc_type, content, method, pages
                        )
        except zipfile.BadZipFile:
            print(f"Warning: Not a valid ZIP file for tender {tender_id}")

    def _extract_single(
        self, filename: str, file_bytes: io.BytesIO
    ) -> Tuple[str, str, int]:
        """Extract text from a single file."""
        filename_lower = filename.lower()
        file_bytes.seek(0)

        try:
            if filename_lower.endswith(".pdf"):
                return self._extract_pdf(file_bytes)
            elif filename_lower.endswith(".docx"):
                return self._extract_docx(file_bytes)
            elif filename_lower.endswith(".doc"):
                # Legacy DOC format - try as DOCX first
                try:
                    return self._extract_docx(file_bytes)
                except:
                    return ("", "unsupported", 0)
            elif filename_lower.endswith((".xlsx", ".xls")):
                return self._extract_excel(file_bytes)
            else:
                return ("", "unsupported", 0)
        except Exception as e:
            print(f"Extraction error for {filename}: {e}")
            return ("", "error", 0)

    def _extract_pdf(self, file_bytes: io.BytesIO) -> Tuple[str, str, int]:
        """Extract text from PDF. Use OCR if needed."""
        try:
            reader = PdfReader(file_bytes)
            text_parts = []
            page_count = len(reader.pages)

            for page in reader.pages:
                text = page.extract_text() or ""
                text_parts.append(text)

            full_text = "\n\n".join(text_parts)

            # Check if we got meaningful text
            if len(full_text.strip()) < 100:
                # Likely scanned PDF - use OCR
                return self._ocr_pdf(file_bytes, page_count)

            return (full_text, "pypdf", page_count)

        except Exception as e:
            print(f"PDF extraction error: {e}")
            return ("", "error", 0)

    def _ocr_pdf(self, file_bytes: io.BytesIO, page_count: int) -> Tuple[str, str, int]:
        """OCR a scanned PDF using PaddleOCR."""
        try:
            # Lazy load PaddleOCR
            if self.ocr_engine is None:
                from paddleocr import PaddleOCR
                self.ocr_engine = PaddleOCR(use_angle_cls=True, lang="fr", use_gpu=False)

            # Convert PDF pages to images and OCR
            # This is a simplified version - full implementation would use pdf2image
            file_bytes.seek(0)
            
            # For now, return empty with OCR method marker
            # Full implementation would convert PDF to images first
            return ("", "paddleocr_pending", page_count)

        except Exception as e:
            print(f"OCR error: {e}")
            return ("", "ocr_error", page_count)

    def _extract_docx(self, file_bytes: io.BytesIO) -> Tuple[str, str, int]:
        """Extract text from DOCX file."""
        try:
            doc = DocxDocument(file_bytes)
            text_parts = [para.text for para in doc.paragraphs if para.text.strip()]
            return ("\n".join(text_parts), "docx", len(doc.paragraphs))
        except Exception as e:
            print(f"DOCX extraction error: {e}")
            return ("", "error", 0)

    def _extract_excel(self, file_bytes: io.BytesIO) -> Tuple[str, str, int]:
        """Extract text from Excel file."""
        try:
            df = pd.read_excel(file_bytes, sheet_name=None)
            text_parts = []
            sheet_count = 0

            for sheet_name, sheet_df in df.items():
                sheet_count += 1
                text_parts.append(f"=== {sheet_name} ===")
                text_parts.append(sheet_df.to_string())

            return ("\n\n".join(text_parts), "xlsx", sheet_count)
        except Exception as e:
            print(f"Excel extraction error: {e}")
            return ("", "error", 0)

    def _classify_document(self, text: str) -> str:
        """Classify document type based on content."""
        text_lower = text.lower()[:2000]  # Check first 2000 chars

        for doc_type, keywords in CLASSIFICATION_KEYWORDS.items():
            for keyword in keywords:
                if keyword in text_lower:
                    return doc_type

        return "OTHER"

    async def _store_document(
        self,
        tender_id: str,
        filename: str,
        doc_type: str,
        content: str,
        method: str,
        pages: int,
    ):
        """Store extracted document in database."""
        if not supabase:
            print(f"Warning: No database. Document {filename} not stored.")
            return

        await supabase.insert("tender_documents", {
            "tender_id": tender_id,
            "document_type": doc_type,
            "original_filename": filename,
            "page_count": pages,
            "extracted_text": content[:50000] if content else None,  # Limit size
            "extraction_method": method,
        })
